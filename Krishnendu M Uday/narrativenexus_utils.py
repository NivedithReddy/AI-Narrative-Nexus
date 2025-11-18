import os
import csv
from typing import List

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


def ensure_data_dir(data_dir: str = "sample_data") -> str:
    """Ensure the data directory exists and return its path."""
    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
    return data_dir


class _DummyUploaded:
    """Internal helper used for tests or compatibility if needed."""

    def __init__(self, name: str, data: bytes):
        self.name = name
        self._data = data

    def getbuffer(self) -> bytes:
        return self._data


def save_uploaded_file(uploaded_file, data_dir: str = "sample_data") -> str:
    """Save a Streamlit uploaded file-like object to disk and return path."""
    ensure_data_dir(data_dir)
    file_path = os.path.join(data_dir, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return file_path


def get_sample_files(data_dir: str = "sample_data") -> List[str]:
    ensure_data_dir(data_dir)
    return sorted(os.listdir(data_dir))


def parse_preview(file_path: str, max_lines: int = 20) -> str:
    """Return a text preview for supported file types (.txt, .csv, .docx).

    - For .txt: return first `max_lines` lines.
    - For .csv: return first `max_lines` rows as comma-separated strings.
    - For .docx: return first `max_lines` paragraphs (if python-docx installed).
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                lines = []
                for i, line in enumerate(f):
                    if i >= max_lines:
                        break
                    lines.append(line.rstrip("\n"))
            return "\n".join(lines)
        except Exception as e:
            return f"[Error reading text file: {e}]"

    if ext == ".csv":
        try:
            with open(file_path, newline="", encoding="utf-8", errors="replace") as csvfile:
                reader = csv.reader(csvfile)
                lines = []
                for i, row in enumerate(reader):
                    if i >= max_lines:
                        break
                    # join with comma+space for nicer preview
                    lines.append(", ".join(row))
            return "\n".join(lines)
        except Exception as e:
            return f"[Error reading CSV file: {e}]"

    if ext == ".docx":
        if not DOCX_AVAILABLE:
            return "[python-docx not installed — cannot preview .docx files]"
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs[:max_lines]]
            return "\n".join(paragraphs)
        except Exception as e:
            return f"[Error reading DOCX file: {e}]"

    return f"[Preview not available for {ext} files]"


def read_full_text(file_path: str) -> str:
    """Read and return the full text content from a file.
    
    Supported formats: .txt, .csv, .docx
    - For .txt: returns all text content
    - For .csv: joins all cells with spaces (good for text analysis)
    - For .docx: extracts all paragraph text
    
    This function performs basic cleaning:
    - Replaces multiple whitespace/newlines with single spaces
    - Strips leading/trailing whitespace
    - Handles encoding errors gracefully
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            # Basic cleaning: normalize whitespace
            text = " ".join(text.split())
            return text
        except Exception as e:
            return f"[Error reading text file: {e}]"

    if ext == ".csv":
        try:
            with open(file_path, newline="", encoding="utf-8", errors="replace") as csvfile:
                reader = csv.reader(csvfile)
                all_text = []
                for row in reader:
                    # Join all cells in the row with space
                    all_text.append(" ".join(row))
            # Join all rows with space and normalize whitespace
            text = " ".join(all_text)
            text = " ".join(text.split())
            return text
        except Exception as e:
            return f"[Error reading CSV file: {e}]"

    if ext == ".docx":
        if not DOCX_AVAILABLE:
            return "[python-docx not installed — cannot read .docx files]"
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs]
            text = " ".join(paragraphs)
            # Normalize whitespace
            text = " ".join(text.split())
            return text
        except Exception as e:
            return f"[Error reading DOCX file: {e}]"

    return f"[Cannot read {ext} files — unsupported format]"
